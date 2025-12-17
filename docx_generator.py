# docx_generator.py
from __future__ import annotations

from datetime import datetime
from typing import Any, Dict, Optional

try:
    from docx import Document  # package "python-docx"
except Exception as e:
    # Message clair si python-docx n'est pas embarqué
    raise RuntimeError(
        "Le module 'python-docx' est requis pour générer des fichiers .docx. "
        "Ajoute 'python-docx' à requirements.txt et rebuild."
    ) from e


def generate_docx(data: Any, output_path: str) -> str:
    """
    Génère un DOCX simple à partir de data (dict/str/whatever).
    Retourne le chemin du fichier généré.
    """
    doc = Document()
    doc.add_heading("Response Report Generator", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().isoformat(timespec='seconds')}")

    doc.add_paragraph("")  # spacer

    if isinstance(data, dict):
        for k, v in data.items():
            doc.add_heading(str(k), level=2)
            doc.add_paragraph(str(v))
    else:
        doc.add_paragraph(str(data))

    doc.save(output_path)
    return output_path


# Alias “au cas où” ton app importait un autre nom
create_docx = generate_docx
build_docx = generate_docx
